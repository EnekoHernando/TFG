import os
import logging
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from llama_index.core import Document, VectorStoreIndex, StorageContext
from llama_index.core.node_parser import SemanticSplitterNodeParser
from llama_index.embeddings.openai import OpenAIEmbedding
from langchain_openai import ChatOpenAI
from nltk.tokenize import sent_tokenize

# Asegurar que el tokenizador de NLTK esté listo para usar
import nltk
nltk.download('punkt')

# Configuración de logging
logging.basicConfig(filename='index_creation.log', level=logging.INFO)
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAIKEY")

def extract_text_from_documents(directory_path):
    documents = []
    for idx, filename in enumerate(os.listdir(directory_path)):
        file_path = os.path.join(directory_path, filename)
        file_metadata = {"filename": filename}
        try:
            if filename.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            elif filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            document = Document(text=text, metadata=file_metadata)
            documents.append(document)
        except Exception as e:
            logging.error(f"Failed to process {filename}: {str(e)}")
    return documents

def extract_text_from_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() if page.extract_text() else ""
    return text

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])

def construct_index(directory_path, index_path):
    documents = extract_text_from_documents(directory_path)
    embed_model = OpenAIEmbedding(api_key=os.environ["OPENAI_API_KEY"])
    llm_predictor = ChatOpenAI(temperature=0.7, model_name="gpt-3.5-turbo", openai_api_key=os.environ["OPENAI_API_KEY"])
    splitter = SemanticSplitterNodeParser(
        embed_model=embed_model,
        buffer_size=1,
        sentence_splitter=sent_tokenize,
        include_metadata=True,
        include_prev_next_rel=True
    )
    nodes = splitter.build_semantic_nodes_from_documents(documents, show_progress=False)
    index = VectorStoreIndex(nodes, llm_predictor=llm_predictor)
    index.storage_context.persist(index_path)
    logging.info("Índice construido y guardado correctamente.")
    return index

def delete_index(index_path):
    # Eliminar el índice existente
    for root, dirs, files in os.walk(index_path, topdown=False):
        for name in files:
            os.remove(os.path.join(root, name))
        for name in dirs:
            os.rmdir(os.path.join(root, name))
    os.rmdir(index_path)
    logging.info(f"Índice '{index_path}' eliminado correctamente.")
    print(f"Índice '{index_path}' eliminado correctamente.")

def main():
    indices_dir = "indices"
    if not os.path.exists(indices_dir):
        os.makedirs(indices_dir)

    while True:
        choice = input("¿Desea crear un nuevo índice, actualizar uno existente, borrar uno existente o salir? (nuevo/actualizar/borrar/exit): ").strip().lower()
        
        if choice == "exit":
            logging.info("Saliendo del programa.")
            print("Saliendo del programa.")
            break
        
        if choice == "nuevo":
            index_name = input("Ingrese el nombre para el nuevo índice: ").strip()
            index_path = os.path.join(indices_dir, index_name)
            if os.path.exists(index_path):
                overwrite = input(f"El índice '{index_name}' ya existe. ¿Desea sobrescribirlo? (sí/no): ").strip().lower()
                if overwrite != "sí":
                    logging.info(f"Operación cancelada. El índice '{index_name}' ya existe y no se sobrescribió.")
                    print("Operación cancelada.")
                    continue
                else:
                    delete_index(index_path)
            construct_index("docs", index_path)
            logging.info(f"Índice '{index_name}' creado y guardado correctamente.")
            print(f"Índice '{index_name}' creado y guardado correctamente.")
        
        elif choice == "actualizar":
            existing_indices = [f for f in os.listdir(indices_dir) if os.path.isdir(os.path.join(indices_dir, f))]
            if not existing_indices:
                logging.info("No hay índices existentes para actualizar.")
                print("No hay índices existentes para actualizar.")
                continue
            
            print("Índices existentes:")
            for i, index_name in enumerate(existing_indices):
                print(f"{i + 1}. {index_name}")
            
            try:
                index_choice = int(input("Seleccione el número del índice que desea actualizar: ").strip())
                if 1 <= index_choice <= len(existing_indices):
                    index_name = existing_indices[index_choice - 1]
                    index_path = os.path.join(indices_dir, index_name)
                    delete_index(index_path)
                    # Crear el nuevo índice
                    construct_index("docs", index_path)
                    logging.info(f"Índice '{index_name}' actualizado correctamente.")
                    print(f"Índice '{index_name}' actualizado correctamente.")
                else:
                    logging.error("Selección inválida.")
                    print("Selección inválida.")
            except ValueError:
                logging.error("Entrada inválida, se esperaba un número.")
                print("Entrada inválida, se esperaba un número.")
        
        elif choice == "borrar":
            existing_indices = [f for f in os.listdir(indices_dir) if os.path.isdir(os.path.join(indices_dir, f))]
            if not existing_indices:
                logging.info("No hay índices existentes para borrar.")
                print("No hay índices existentes para borrar.")
                continue
            
            print("Índices existentes:")
            for i, index_name in enumerate(existing_indices):
                print(f"{i + 1}. {index_name}")
            
            try:
                index_choice = int(input("Seleccione el número del índice que desea borrar: ").strip())
                if 1 <= index_choice <= len(existing_indices):
                    index_name = existing_indices[index_choice - 1]
                    index_path = os.path.join(indices_dir, index_name)
                    delete_index(index_path)
                    logging.info(f"Índice '{index_name}' borrado correctamente.")
                    print(f"Índice '{index_name}' borrado correctamente.")
                else:
                    logging.error("Selección inválida.")
                    print("Selección inválida.")
            except ValueError:
                logging.error("Entrada inválida, se esperaba un número.")
                print("Entrada inválida, se esperaba un número.")
        
        else:
            logging.error("Opción no reconocida. Por favor, elija 'nuevo', 'actualizar', 'borrar' o 'exit'.")
            print("Opción no reconocida. Por favor, elija 'nuevo', 'actualizar', 'borrar' o 'exit'.")

if __name__ == "__main__":
    main()