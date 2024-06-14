import os
import logging
import threading
from tkinter import Tk, Toplevel, filedialog, messagebox, Button, Label, Entry, Listbox, Scrollbar, VERTICAL, END, Text, SINGLE
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from llama_index.core import Document, VectorStoreIndex
from llama_index.core.node_parser import SemanticSplitterNodeParser
from llama_index.embeddings.openai import OpenAIEmbedding
from langchain_openai import ChatOpenAI
from nltk.tokenize import sent_tokenize

# Asegurar que el tokenizador de NLTK esté listo para usar
import nltk
nltk.download('punkt')

# Configuración de logging
logging.basicConfig(filename='index_creation.log', level=logging.INFO)
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAIKEY")

file_paths = []
log_update_interval = 1000  # Intervalo de actualización de logs en milisegundos

def clear_logs():
    open('index_creation.log', 'w').close()

def extract_text_from_documents(file_paths):
    documents = []
    for file_path in file_paths:
        filename = os.path.basename(file_path)
        file_metadata = {"filename": filename}
        try:
            if filename.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            elif filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            document = Document(text=text, metadata=file_metadata)
            documents.append(document)
        except Exception as e:
            logging.error(f"Failed to process {filename}: {str(e)}")
    return documents

def extract_text_from_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() if page.extract_text() else ""
    return text

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])

def construct_index(file_paths, index_path):
    clear_logs()
    documents = extract_text_from_documents(file_paths)
    embed_model = OpenAIEmbedding(api_key=os.environ["OPENAI_API_KEY"])
    llm_predictor = ChatOpenAI(temperature=0.7, model_name="gpt-3.5-turbo", openai_api_key=os.environ["OPENAI_API_KEY"])
    splitter = SemanticSplitterNodeParser(
        embed_model=embed_model,
        buffer_size=1,
        sentence_splitter=sent_tokenize,
        include_metadata=True,
        include_prev_next_rel=True
    )
    nodes = splitter.build_semantic_nodes_from_documents(documents, show_progress=False)
    index = VectorStoreIndex(nodes, llm_predictor=llm_predictor)
    index.storage_context.persist(index_path)
    logging.info("Índice construido y guardado correctamente.")

def delete_index(index_path):
    clear_logs()
    for root, dirs, files in os.walk(index_path, topdown=False):
        for name in files:
            os.remove(os.path.join(root, name))
        for name in dirs:
            os.rmdir(os.path.join(root, name))
    os.rmdir(index_path)
    logging.info(f"Índice '{index_path}' eliminado correctamente.")
    messagebox.showinfo("Info", f"Índice '{index_path}' eliminado correctamente.")

def add_files():
    new_files = filedialog.askopenfilenames(title="Seleccione archivos", filetypes=[("Documentos", "*.pdf *.docx *.txt")])
    for file in new_files:
        if file not in file_paths:
            file_paths.append(file)
            file_listbox.insert(END, os.path.basename(file))

def run_in_thread(target):
    thread = threading.Thread(target=target)
    thread.start()

def create_index():
    index_name = index_name_entry.get().strip()
    if not index_name:
        messagebox.showerror("Error", "El nombre del índice no puede estar vacío.")
        return

    if not file_paths:
        messagebox.showerror("Error", "No se seleccionaron archivos.")
        return

    run_in_thread(lambda: create_index_task(index_name))

def create_index_task(index_name):
    indices_dir = "indices"
    if not os.path.exists(indices_dir):
        os.makedirs(indices_dir)

    index_path = os.path.join(indices_dir, index_name)
    if os.path.exists(index_path):
        overwrite = messagebox.askyesno("Sobrescribir", f"El índice '{index_name}' ya existe. ¿Desea sobrescribirlo?")
        if not overwrite:
            return
        else:
            delete_index(index_path)

    construct_index(file_paths, index_path)
    messagebox.showinfo("Info", f"Índice '{index_name}' creado y guardado correctamente.")
    file_paths.clear()
    file_listbox.delete(0, END)
    display_logs()

def update_index():
    existing_indices = [f for f in os.listdir("indices") if os.path.isdir(os.path.join("indices", f))]
    if not existing_indices:
        messagebox.showerror("Error", "No hay índices existentes para actualizar.")
        return

    index_name = index_name_entry.get().strip()
    if index_name not in existing_indices:
        messagebox.showerror("Error", f"El índice '{index_name}' no existe.")
        return

    if not file_paths:
        messagebox.showerror("Error", "No se seleccionaron archivos.")
        return

    run_in_thread(lambda: update_index_task(index_name))

def update_index_task(index_name):
    index_path = os.path.join("indices", index_name)
    delete_index(index_path)
    construct_index(file_paths, index_path)
    messagebox.showinfo("Info", f"Índice '{index_name}' actualizado correctamente.")
    file_paths.clear()
    file_listbox.delete(0, END)
    display_logs()

def delete_index_gui():
    existing_indices = [f for f in os.listdir("indices") if os.path.isdir(os.path.join("indices", f))]
    if not existing_indices:
        messagebox.showerror("Error", "No hay índices existentes para borrar.")
        return

    delete_window = Toplevel(root)
    delete_window.title("Seleccionar Índice para Borrar")

    Label(delete_window, text="Seleccione un índice para borrar:").pack(padx=10, pady=10)
    
    index_listbox = Listbox(delete_window, selectmode=SINGLE, width=50, height=10)
    for idx in existing_indices:
        index_listbox.insert(END, idx)
    index_listbox.pack(padx=10, pady=10)

    Button(delete_window, text="Borrar", command=lambda: confirm_delete_index(index_listbox, delete_window)).pack(padx=10, pady=10)

def confirm_delete_index(index_listbox, window):
    selected_index = index_listbox.curselection()
    if not selected_index:
        messagebox.showerror("Error", "No se seleccionó ningún índice.")
        return
    index_name = index_listbox.get(selected_index)
    run_in_thread(lambda: delete_index_task(index_name))
    window.destroy()

def delete_index_task(index_name):
    index_path = os.path.join("indices", index_name)
    delete_index(index_path)
    display_logs()

def display_logs():
    log_text.config(state='normal')
    log_text.delete(1.0, END)
    with open('index_creation.log', 'r') as log_file:
        log_text.insert(END, log_file.read())
    log_text.config(state='disabled')

def update_log_periodically():
    display_logs()
    root.after(log_update_interval, update_log_periodically)

# Configuración de la interfaz gráfica
root = Tk()
root.title("Gestor de Índices")

Label(root, text="Nombre del Índice:").grid(row=0, column=0, padx=10, pady=10)
index_name_entry = Entry(root)
index_name_entry.grid(row=0, column=1, padx=10, pady=10, columnspan=2)

Button(root, text="Añadir Archivos", command=add_files).grid(row=1, column=0, padx=10, pady=10, columnspan=3)

file_listbox = Listbox(root, selectmode="multiple", width=100, height=10)
file_listbox.grid(row=2, column=0, padx=10, pady=10, columnspan=3)

scrollbar = Scrollbar(root, orient=VERTICAL, command=file_listbox.yview)
scrollbar.grid(row=2, column=3, sticky='ns')
file_listbox.config(yscrollcommand=scrollbar.set)

Button(root, text="Crear Índice", command=create_index).grid(row=3, column=0, padx=10, pady=10)
Button(root, text="Actualizar Índice", command=update_index).grid(row=3, column=1, padx=10, pady=10)
Button(root, text="Borrar Índice", command=delete_index_gui).grid(row=3, column=2, padx=10, pady=10)

Label(root, text="Logs:").grid(row=4, column=0, padx=10, pady=10)
log_text = Text(root, state='disabled', width=100, height=10)
log_text.grid(row=5, column=0, padx=10, pady=10, columnspan=3)

scrollbar_log = Scrollbar(root, orient=VERTICAL, command=log_text.yview)
scrollbar_log.grid(row=5, column=3, sticky='ns')
log_text.config(yscrollcommand=scrollbar_log.set)

# Iniciar actualización periódica de logs
update_log_periodically()

root.mainloop()