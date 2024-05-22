import os
import logging
import socket
import time
from PyPDF2 import PdfReader
from docx import Document as DocxDocument  # To avoid naming conflict with our Document class
from llama_index.core import Document, VectorStoreIndex, StorageContext, load_index_from_storage
from llama_index.core.node_parser import SemanticSplitterNodeParser
from llama_index.embeddings.openai import OpenAIEmbedding
from langchain_openai import ChatOpenAI  # Updated import based on deprecation notice
from nltk.tokenize import sent_tokenize

# Asegurar que el tokenizador de NLTK esté listo para usar
import nltk
nltk.download('punkt')

# Configuración de logging para registrar eventos importantes
logging.basicConfig(filename='server.log', level=logging.INFO)
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAIKEY")

def extract_text_from_documents(directory_path):
    """
    Extrae texto de los documentos en un directorio dado.
    Soporta archivos PDF, DOCX y TXT.
    """
    print(f"Extrayendo texto de los documentos")
    documents = []
    for idx, filename in enumerate(os.listdir(directory_path)):
        file_path = os.path.join(directory_path, filename)
        file_metadata = {"filename": filename}
        print(f"Documento actual: {filename}")
        try:
            # Determinar el tipo de archivo y extraer texto en consecuencia
            if filename.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            elif filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            # Crear el objeto Document con el texto extraído y metadatos
            document = Document(text=text, metadata=file_metadata)
            print("Documento terminado... \n")
            documents.append(document)
        except Exception as e:
            # Registrar cualquier error durante el procesamiento del archivo
            logging.error(f"Failed to process {filename}: {str(e)}")
    return documents

def extract_text_from_pdf(file_path):
    """
    Extrae texto de un archivo PDF.
    """
    text = ''
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            # Extraer texto de cada página del PDF
            text += page.extract_text() if page.extract_text() else ""
    return text

def extract_text_from_docx(file_path):
    """
    Extrae texto de un archivo DOCX.
    """
    doc = DocxDocument(file_path)
    # Unir el texto de todos los párrafos del documento
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])

def construct_index(directory_path, index_path):
    """
    Construye un índice a partir de documentos en un directorio dado y lo guarda en index_path.
    """
    print(f"Nuevo Indice con nombre: {index_path}; siendo generado a partir de la documentación en: {directory_path}")
    # Extraer texto y crear objetos Document
    documents = extract_text_from_documents(directory_path)
    
    # Inicializar los modelos de embedding y LLM
    embed_model = OpenAIEmbedding(api_key=os.environ["OPENAI_API_KEY"])
    llm_predictor = ChatOpenAI(temperature=0.7, model_name="gpt-3.5-turbo", openai_api_key=os.environ["OPENAI_API_KEY"])
    
    # Inicializar el SemanticSplitterNodeParser con el modelo de embedding
    splitter = SemanticSplitterNodeParser(
        embed_model=embed_model,
        buffer_size=1,  # Número de oraciones a agrupar al evaluar la similitud semántica
        sentence_splitter=sent_tokenize,  # Usando el tokenizador de oraciones de NLTK
        include_metadata=True,
        include_prev_next_rel=True
    )
    
    print("Convirtiendo texto extraído en nodos")
    # Convertir el texto extraído en nodos usando el parser de nodos
    nodes = splitter.build_semantic_nodes_from_documents(documents, show_progress=False)
    
    print("Creando índices a partir de nodos")
    # Crear el índice usando los nodos y almacenarlo
    index = VectorStoreIndex(nodes, llm_predictor=llm_predictor)
    # Persistir el índice
    index.storage_context.persist(index_path)
    logging.info("Índice construido y guardado correctamente.")
    print("Índice creado y almacenado")
    return index

def chatbot(input_text, index_path):
    """
    Maneja las consultas al chatbot utilizando un índice preexistente.
    """
    logging.info(f"Recibido texto: {input_text}")
    # Reconstruir el contexto de almacenamiento
    storage_context = StorageContext.from_defaults(persist_dir=index_path)
    # Cargar índice desde el contexto de almacenamiento
    new_index = load_index_from_storage(storage_context)
    # Crear un motor de consulta a partir del índice cargado
    new_query_engine = new_index.as_query_engine()
    # Realizar la consulta y obtener la respuesta
    response = new_query_engine.query(input_text)
    logging.info(f"Texto de respuesta: {response}")
    return response.response

# Crear el socket del servidor sin SSL
server_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
host = 'localhost'
port = 12345
server_socket.bind((host, port))
server_socket.listen()

# Crear la carpeta de índices si no existe
indices_dir = "indices"
if not os.path.exists(indices_dir):
    os.makedirs(indices_dir)

try:
    # Preguntar al usuario si desea usar un índice creado o generar uno nuevo
    message = input("¿Quieres usar el índice creado o generar uno nuevo? (creado/generar): ")
    if message == "generar":
        name = input("¿Cómo quieres llamarlo?: ")
        # Crear el índice en la carpeta de índices
        index_path = os.path.join(indices_dir, name)
        construct_index("docs", index_path)
    else:
        # Listar los índices disponibles
        indices = [d for d in os.listdir(indices_dir) if os.path.isdir(os.path.join(indices_dir, d))]
        print("Índices disponibles:")
        for idx in indices:
            print(f" - {idx}")
        name = input("¿Cuál quieres usar?: ")
        index_path = os.path.join(indices_dir, name)
except Exception as e:
    logging.info(f"Servidor cerrado por {e}.")
finally:
    try:
        logging.info('Ya se puede conectar al servidor')
        print("Ya se puede conectar al servidor")
        while True:
            client_socket, addr = server_socket.accept()
            print(f"Conectado con {addr}")
            logging.info(f'Conectado con {addr}')

            try:
                while True:
                    # Recibir texto del cliente
                    input_text = client_socket.recv(1024).decode('utf-8')
                    print(f'Texto recibido: {input_text}')
                    if input_text == "bye":
                        break
                    
                    # Calcular el tiempo de respuesta
                    start_time = time.time()
                    # Obtener respuesta del chatbot y enviarla de vuelta al cliente
                    response = chatbot(input_text, index_path)
                    end_time = time.time()
                    response_time = end_time - start_time
                    logging.info(f'Tiempo de respuesta: {response_time} segundos')
                    
                    client_socket.sendall(response.encode('utf-8'))
            except Exception as e:
                logging.error(f"Error al manejar la conexión del cliente {addr}: {e}")
            finally:
                client_socket.close()
                logging.info(f"Conexión con el cliente {addr} cerrada.")
    except KeyboardInterrupt:
        logging.info("Servidor cerrado por interrupción del teclado.")
    finally:
        server_socket.close()
        logging.info("Socket del servidor cerrado.")